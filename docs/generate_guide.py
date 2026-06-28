"""Generates the layman's technical explanation guide for CyberGuard AI."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUT_PATH = "CyberGuard_AI_Technical_Guide.docx"


def _r(para, text, bold=False, italic=False, size=12, mono=False, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Courier New" if mono else "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 112, 192)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 70, 127)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(12)
    return p


def mixed(doc, parts):
    """parts = list of (text, bold, mono) tuples."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    for text, bold, mono in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Courier New" if mono else "Calibri"
        r.font.size = Pt(11 if mono else 12)
    return p


def bul(doc, items):
    for item in items:
        if isinstance(item, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            for seg, bold, mono in item:
                r = p.add_run(seg)
                r.bold = bold
                r.font.name = "Courier New" if mono else "Calibri"
                r.font.size = Pt(11 if mono else 12)
        else:
            p = doc.add_paragraph(item, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(12)


def code_block(doc, code, caption=None):
    if caption:
        h3(doc, caption)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(code)
    r.font.name = "Courier New"
    r.font.size = Pt(9)


def qa(doc, question, answer):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.25)
    r = p.add_run("Q: " + question)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(192, 0, 0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf2 = p2.paragraph_format
    pf2.space_after = Pt(4)
    pf2.left_indent = Inches(0.25)
    pf2.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf2.line_spacing = 1.5
    r2 = p2.add_run("A: " + answer)
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.0)

    # ── COVER ────────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CYBERGUARD AI")
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0, 112, 192)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Web-Based Cybersecurity Advisory & Awareness System")
    r.font.name = "Calibri"; r.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TECHNICAL EXPLANATION GUIDE")
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(18)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A complete, plain-English walkthrough of every component:\n"
        "what it does, how it works, why it was built that way.\n\n"
        "Use this document to explain every part of your project\n"
        "to your supervisor, examiner, or anyone who asks."
    )
    r.font.name = "Calibri"; r.font.size = Pt(13); r.italic = True
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1 — WHAT IS THIS PROJECT?
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 1: WHAT IS THIS PROJECT, AND WHY DOES IT EXIST?")

    h2(doc, "1.1  The Big Picture (In One Paragraph)")
    body(doc,
         "CyberGuard AI is a website that gives people cybersecurity advice — like having "
         "a professional security expert available 24/7 in your browser for free. It is "
         "powered by Meta's Llama 3.1 AI model (the same family of AI technology "
         "behind many modern AI assistants), accessed through a service called Hugging "
         "Face. A user opens the site, types a cybersecurity question, and gets an "
         "intelligent, expert-quality answer within seconds.")

    h2(doc, "1.2  The Problem It Solves")
    body(doc,
         "Imagine you receive a suspicious email and want to know if it is a phishing "
         "scam. Or you want to create a strong password but don't know what makes one "
         "strong. Or your company wants to train staff on cybersecurity but cannot "
         "afford a consultant. These are everyday problems that affect millions of "
         "people — especially in Nigeria, where affordable cybersecurity advisory is "
         "scarce.")
    bul(doc, [
        "Professional cybersecurity consultants are expensive.",
        "Free online resources are scattered, technical, and hard to understand.",
        "Most free tools do only one thing (e.g., check a password OR scan a URL, "
        "but never both, and never with conversational explanation).",
        "There is no AI-powered, free, browser-accessible tool that combines advice, "
        "education, quizzes, password checking, and URL analysis in one place.",
    ])
    body(doc,
         "CyberGuard AI solves all of these problems in a single website.")

    h2(doc, "1.3  Who Is It For?")
    bul(doc, [
        "Students wanting to learn cybersecurity concepts interactively.",
        "Office workers who want to check whether a suspicious email or link is safe.",
        "Small business owners who cannot afford a dedicated security team.",
        "Lecturers and trainers who want a demonstration tool for security awareness "
        "sessions.",
        "Anyone who wants to create stronger passwords or understand news about "
        "software vulnerabilities (CVEs).",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2 — HOW THE SYSTEM IS STRUCTURED
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 2: HOW THE SYSTEM IS STRUCTURED")

    h2(doc, "2.1  The Three-Tier Architecture (Simple Analogy)")
    body(doc,
         "Think of CyberGuard AI like a restaurant:")
    bul(doc, [
        "The Customer (Frontend) — sits at a table, reads the menu, places an order. "
        "In our system, this is the web page the user sees and interacts with in their "
        "browser.",
        "The Waiter (Backend/Flask Server) — takes the customer's order to the kitchen, "
        "brings back the food. In our system, this is the Python Flask application "
        "running on a computer that receives requests from the user and sends them to "
        "the AI.",
        "The Kitchen/Chef (Hugging Face AI) — actually prepares the food. In our "
        "system, this is Meta's Llama 3.1 AI model running on Hugging Face's servers "
        "that generates the intelligent responses.",
    ])

    h2(doc, "2.2  The Three Files That Make Everything Work")
    body(doc,
         "The entire system lives in three core files plus supporting configuration:")

    h3(doc, "app.py  (the Brain — Backend Python Server)")
    body(doc,
         "This is the heart of the system. It is a Python program that runs continuously "
         "on your computer (or a server). When a user's browser sends a request — for "
         "example, 'answer this cybersecurity question' — app.py receives it, constructs "
         "the right AI prompt, calls the Hugging Face API, gets the AI's answer, and "
         "sends it back to the browser as JSON data.")

    h3(doc, "templates/index.html  (the Face — What the User Sees)")
    body(doc,
         "This is the HTML file that defines the structure of the web page — all the "
         "buttons, panels, input boxes, and layout. Flask serves this file when a user "
         "visits the site. It is a Single-Page Application (SPA), meaning the entire "
         "site loads once and navigation between features happens without page reloads.")

    h3(doc, "static/js/app.js  (the Intelligence — Browser Logic)")
    body(doc,
         "This is the JavaScript file that runs inside the user's browser and makes the "
         "site interactive. It handles: tab switching, sending requests to the backend, "
         "displaying AI responses, the password calculator, the quiz scoring, the matrix "
         "rain animation, and the connection status checker.")

    h3(doc, "Supporting Files")
    bul(doc, [
        "static/css/style.css — All the visual styling: dark colours, neon green "
        "accents, sidebar layout, animations.",
        ".env — Stores secret configuration values (API key, Flask secret key). "
        "Never committed to GitHub.",
        "requirements.txt — Lists all Python libraries needed to run the project.",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3 — THE BACKEND (app.py) EXPLAINED IN FULL
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 3: THE BACKEND — app.py EXPLAINED LINE BY LINE")

    h2(doc, "3.1  What is Flask and Why Was It Chosen?")
    body(doc,
         "Flask is a Python web framework — a toolkit that makes it easy to build "
         "websites and APIs in Python. When you run a Flask app, it starts a small "
         "web server on your computer that listens for incoming HTTP requests (the same "
         "type of requests your browser makes every time you visit a website).")
    body(doc,
         "Flask was chosen because:")
    bul(doc, [
        "It is lightweight and easy to understand — perfect for a final year project "
        "where the focus is on the AI features, not web framework complexity.",
        "It requires minimal configuration to get a web server running.",
        "It works perfectly with Python 3.6.8, which was the Python version available "
        "on the development machine.",
        "It has a large community and extensive documentation.",
        "It natively supports returning JSON responses, which is exactly what the "
        "frontend needs.",
    ])

    h2(doc, "3.2  The Imports and Configuration Block")
    code_block(doc,
               "import os, json, re\n"
               "from datetime import datetime\n"
               "import requests\n"
               "from flask import Flask, render_template, request, jsonify\n"
               "from flask_cors import CORS\n"
               "from dotenv import load_dotenv\n\n"
               "load_dotenv()\n"
               "app = Flask(__name__)\n"
               "CORS(app)")
    body(doc,
         "What each import does:")
    bul(doc, [
        "os — reads environment variables (the API key stored in the .env file).",
        "json — parses and creates JSON data (the format used for API communication).",
        "re — regular expressions, used to extract JSON from AI responses that may "
        "contain extra text.",
        "datetime — gets the current date and time, used for daily tip timestamps.",
        "requests — sends HTTP requests to the Hugging Face API (like a browser, "
        "but from Python).",
        "Flask, render_template, request, jsonify — Flask components for running the "
        "server, serving HTML, reading incoming request data, and sending JSON responses.",
        "CORS — Cross-Origin Resource Sharing. Without this, the browser would refuse "
        "to let the frontend JavaScript talk to the backend. CORS tells the browser "
        "it is safe to do so.",
        "load_dotenv() — reads the .env file and makes the API key available as an "
        "environment variable.",
    ])

    h2(doc, "3.3  The System Prompt — The AI's Personality")
    code_block(doc,
               'SYSTEM_PROMPT = (\n'
               '    "You are CyberGuard AI, a world-class cybersecurity expert advisor "\n'
               '    "powered by Meta Llama. You specialize in threat analysis, "\n'
               '    "vulnerability assessments, incident response, security best practices..."\n'
               '    "Never assist with actual attacks or help bypass security systems."\n'
               ')')
    body(doc,
         "Think of the system prompt as the 'job description' you give to the AI at "
         "the start of every conversation. Every time we call the AI, we include this "
         "instruction. It tells the AI:")
    bul(doc, [
        "What role to play: a cybersecurity expert named CyberGuard AI.",
        "What topics to focus on: threat analysis, vulnerabilities, best practices, "
        "education.",
        "What to never do: assist with actual attacks or help bypass security.",
        "How to respond: clearly, with focus on defensive measures.",
    ])
    body(doc,
         "This is called 'prompt engineering' — designing the instructions to an AI "
         "model to get the type and quality of output you want. The system prompt is "
         "the most important piece of AI tuning in the project.")

    h2(doc, "3.4  _get_session() — The Proxy-Aware HTTP Client")
    code_block(doc,
               "def _get_session() -> requests.Session:\n"
               "    session = requests.Session()\n"
               "    proxy = os.getenv('HTTPS_PROXY') or os.getenv('HTTP_PROXY')\n"
               "    if proxy:\n"
               "        session.proxies = {'http': proxy, 'https': proxy}\n"
               "    return session")
    body(doc,
         "This function creates an HTTP client (called a 'session') that is used to "
         "send requests to the Hugging Face API. The important thing it does is check "
         "whether proxy settings have been configured.")
    body(doc,
         "Why is this needed? Many university and corporate networks route internet "
         "traffic through a 'proxy server' — an intermediary computer that all "
         "outgoing connections must pass through. If you don't configure the proxy, "
         "any attempt to reach the internet fails with a 'connection refused' or DNS "
         "error. By reading proxy settings from the .env file and applying them to the "
         "session, the system can work on these restricted networks.")

    h2(doc, "3.5  call_llama() — The Core AI Function")
    body(doc,
         "This is the most important function in the entire project. Every single "
         "AI-powered feature calls this one function. Understanding it is understanding "
         "the heart of the system.")

    h3(doc, "What it receives:")
    bul(doc, [
        "messages: A list of conversation messages in the format [{\"role\": \"user\", "
        "\"content\": \"...\"}]. This is the standard format used by AI chat APIs.",
        "max_tokens: The maximum length of the AI's response (measured in 'tokens', "
        "which are roughly 3/4 of a word).",
    ])

    h3(doc, "What it does:")
    bul(doc, [
        "Builds the HTTP request headers including the Authorization header with "
        "the Bearer token (API key).",
        "Builds the JSON payload with the model name, messages, temperature (0.7 — "
        "a balance between creative and factual), and max_tokens.",
        "Calls _get_session() to get the proxy-aware HTTP client.",
        "Sends a POST request to the Hugging Face API endpoint.",
        "Extracts the AI's text from the response: "
        "resp.json()[\"choices\"][0][\"message\"][\"content\"]",
    ])

    h3(doc, "Why temperature=0.7?")
    body(doc,
         "Temperature controls how 'creative' the AI is. At 0.0, it always gives the "
         "most predictable answer. At 1.0, it is very creative and varied but may "
         "produce inconsistent responses. 0.7 is a standard middle ground — responses "
         "are varied enough to feel natural but grounded enough to be accurate for a "
         "technical domain like cybersecurity.")

    h3(doc, "Error Handling — Why Every Error Has Its Own Message:")
    bul(doc, [
        "ConnectionError → 'NETWORK_ERROR:' prefix. Means the machine cannot reach "
        "the internet. Common in university networks. The NETWORK_ERROR: prefix is "
        "detected by the frontend JavaScript to show the red offline banner.",
        "Timeout → Response too slow (> 90 seconds). First API call loads the model "
        "into memory which can take ~30 seconds. User is told to retry.",
        "HTTPError 403 → API key valid but model requires license acceptance on "
        "Hugging Face website.",
        "HTTPError 401 → API key is wrong or expired.",
    ])
    body(doc,
         "Each error gives the user a specific, actionable message rather than a "
         "generic 'something went wrong.' This is good software engineering practice "
         "— it helps the user (or developer) understand exactly what is wrong and "
         "what to do about it.")

    h2(doc, "3.6  The Seven API Endpoints — One by One")

    h3(doc, "Endpoint 1: GET /  (The Home Page)")
    body(doc,
         "The simplest endpoint. When a user types the website address in their "
         "browser, this endpoint responds by serving the index.html file. Flask's "
         "render_template() function reads the HTML file from the templates/ folder "
         "and sends it to the browser.")
    code_block(doc, "@app.route('/')\ndef index():\n    return render_template('index.html')")

    h3(doc, "Endpoint 2: POST /api/chat  (The AI Advisor)")
    body(doc,
         "This is the main conversational endpoint. It accepts:")
    bul(doc, [
        "message: The user's current question.",
        "history: An array of the last few messages in the conversation "
        "(so the AI remembers context).",
    ])
    body(doc,
         "It builds a message array starting with the system prompt, then adds up to "
         "12 previous messages from history, then the new user message. This is how "
         "the AI maintains context — each request includes the full conversation "
         "history. The AI then responds as if it has been part of the conversation "
         "from the beginning.")
    body(doc,
         "Why 12 messages? LLMs have a 'context window' — a maximum amount of text "
         "they can process at once. Limiting history to 12 messages prevents the "
         "prompt from becoming too long and exceeding the model's context window, "
         "which would cause errors.")

    h3(doc, "Endpoint 3: POST /api/threat-analysis  (Threat Deep Dive)")
    body(doc,
         "Receives a threat name (e.g., 'Ransomware') and asks the AI to produce a "
         "structured 5-section analysis:")
    bul(doc, [
        "1. What It Is — definition",
        "2. How It Works — mechanism",
        "3. Warning Signs — detection indicators",
        "4. Prevention Measures — how to stay safe",
        "5. How to Respond if Affected — incident response steps",
    ])
    body(doc,
         "The structure is enforced through the prompt. By explicitly listing the "
         "five sections in the request, the AI consistently formats its response "
         "in the same way regardless of which threat is queried. This is 'structured "
         "output via prompt engineering.'")

    h3(doc, "Endpoint 4: GET /api/quiz  (Quiz Question Generator)")
    body(doc,
         "This is the most technically interesting endpoint because it requires the "
         "AI to produce a response in strict JSON format (not just free text).")
    body(doc,
         "The prompt explicitly tells the AI:")
    code_block(doc,
               "\"Respond ONLY with valid JSON in this exact format:\\n\"\n"
               "{\"question\": \"...\", \"options\": {\"A\": \"...\", \"B\": \"...\", \"C\": \"...\", \"D\": \"...\"},\n"
               "\"correct\": \"A\", \"explanation\": \"...\"}\\n\"\n"
               "\"No extra text, just the JSON object.\"")
    body(doc,
         "Even with this instruction, AI models sometimes add extra text before or "
         "after the JSON. That is why the code uses a regular expression to find and "
         "extract just the JSON portion:")
    code_block(doc,
               "match = re.search(r'\\{[\\s\\S]*\\}', raw)\n"
               "if match:\n"
               "    quiz_data = json.loads(match.group())")
    body(doc,
         "The regex r'\\{[\\s\\S]*\\}' matches everything from the first { to the last } "
         "in the response, extracting the JSON object even if there is surrounding text. "
         "This is a robust pattern for extracting structured data from LLM outputs.")

    h3(doc, "Endpoint 5: GET /api/security-tip  (Daily Security Tip)")
    body(doc,
         "The simplest AI endpoint. Asks for one specific, actionable security tip with "
         "a maximum of 80 words. The constraint on word count is enforced through the "
         "prompt and by setting max_tokens=150 (approximately 100-120 words), ensuring "
         "responses are brief and readable. The tip auto-loads when the user opens the "
         "Daily Tips tab, making the experience feel immediate and responsive.")

    h3(doc, "Endpoint 6: POST /api/analyze-url  (URL Phishing Analyser)")
    body(doc,
         "Receives a URL and instructs the AI to analyse it for phishing indicators "
         "based on structure alone — not by visiting it. This distinction is "
         "important: the system never fetches the URL, which would be dangerous and "
         "could expose the server to malicious content. Instead, the AI analyses "
         "structural features:")
    bul(doc, [
        "Suspicious domain patterns (e.g., paypa1.com instead of paypal.com — "
        "typosquatting)",
        "Unusual TLDs (e.g., .ru, .tk, .xyz in banking contexts)",
        "Excessive subdomains (e.g., login.secure.verify.paypal.com.attackersite.net)",
        "IP addresses instead of domain names in URLs",
        "URL shorteners that hide the real destination",
        "Misleading paths (e.g., https://facebook.com.evil.site/login)",
    ])
    body(doc,
         "The AI returns a Risk Level (Low/Medium/High/Critical) with explanations "
         "of each finding. This gives the user an educational, explanatory analysis "
         "rather than a simple pass/fail score.")

    h3(doc, "Endpoint 7: POST /api/explain-cve  (CVE Vulnerability Lookup)")
    body(doc,
         "A CVE (Common Vulnerabilities and Exposures) is a standardised identifier "
         "for known software vulnerabilities — for example, CVE-2021-44228 is the "
         "identifier for Log4Shell, a critical vulnerability in the Apache Log4j "
         "logging library. The NVD (National Vulnerability Database) publishes "
         "technical CVE descriptions that are often difficult for non-experts to "
         "understand.")
    body(doc,
         "This endpoint accepts a CVE identifier and asks the AI to explain it in "
         "plain language: what software is affected, what an attacker could do, "
         "the CVSS (Common Vulnerability Scoring System) severity score, and what "
         "patches or mitigations are available. The AI's training data includes "
         "extensive documentation of well-known CVEs, making it well-suited to "
         "this translation task.")

    h3(doc, "Endpoint 8: GET /api/status  (Connection Health Check)")
    body(doc,
         "This endpoint does NOT call the AI model. Instead, it makes a lightweight "
         "request to the Hugging Face 'whoami' API to verify that: (1) the API key "
         "is valid, and (2) the server can reach huggingface.co. It returns a JSON "
         "object: {\"ok\": true} or {\"ok\": false, \"reason\": \"...\"}.")
    body(doc,
         "The frontend calls this endpoint on page load and every time the user "
         "clicks Retry. This allows the interface to show the user a clear "
         "'AI Online' or 'AI Offline' status indicator, and to display an "
         "informative red banner when the API is unreachable.")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4 — THE FRONTEND
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 4: THE FRONTEND — WHAT THE USER SEES AND HOW IT WORKS")

    h2(doc, "4.1  Why a Single-Page Application (SPA)?")
    body(doc,
         "A traditional website loads a new HTML page every time you click a link. "
         "A Single-Page Application loads once and then dynamically shows and hides "
         "different sections without ever reloading the page. CyberGuard AI is built "
         "as an SPA for several reasons:")
    bul(doc, [
        "Faster navigation — switching between the Chat tab and the Quiz tab is "
        "instant because no page reload is needed.",
        "Better user experience — the matrix rain animation and connection status "
        "badge persist across tab switches.",
        "Simpler architecture — all features are in one HTML file, making it easy "
        "to maintain.",
    ])

    h2(doc, "4.2  The Visual Design — Why Dark, Why Neon?")
    body(doc,
         "The interface uses a dark colour theme with neon green (#00ff88) and "
         "electric blue (#00d4ff) accents. This aesthetic was deliberately chosen to "
         "match the visual language of professional cybersecurity tools and CTF (Capture "
         "The Flag) platforms — terminal interfaces, security dashboards, and hacker "
         "movies universally use dark backgrounds with bright accent colours. This "
         "immediately communicates to the user that this is a security tool, not a "
         "general-purpose assistant.")
    body(doc,
         "The CSS uses custom properties (variables) for all colours:")
    code_block(doc,
               ":root {\n"
               "  --bg-base: #050b14;   /* very dark blue-black background */\n"
               "  --accent: #00ff88;    /* neon green for highlights */\n"
               "  --cyan: #00d4ff;      /* electric blue for accents */\n"
               "  --warn: #ff9500;      /* orange for Coming Soon badges */\n"
               "}")
    body(doc,
         "Using variables means changing the entire colour scheme requires editing "
         "only these 4 lines — a maintainability best practice.")

    h2(doc, "4.3  The Matrix Rain Animation")
    body(doc,
         "The falling characters animation in the background (like the Matrix movie) "
         "is implemented using the HTML5 Canvas API — a browser feature that allows "
         "drawing graphics with JavaScript.")
    body(doc,
         "How it works:")
    bul(doc, [
        "A <canvas> element is placed behind all other content at z-index: 0 with "
        "very low opacity (4%) so it does not distract from the interface.",
        "An array of 'drops' tracks the Y position of each falling column of text.",
        "Every 60 milliseconds (using setInterval), the canvas is redrawn: first a "
        "semi-transparent black rectangle is drawn over the whole canvas (this creates "
        "the 'fading trail' effect), then a character is drawn at each column's "
        "current position.",
        "Characters are randomly selected from a mix of Japanese katakana, "
        "hexadecimal digits (0-9, A-F), and symbols.",
        "When a drop reaches the bottom of the screen, it resets to the top at "
        "a random position.",
    ])
    body(doc,
         "This is purely cosmetic — it adds atmosphere but has no functional purpose. "
         "It runs in the background using the browser's event loop and does not "
         "interfere with any other functionality.")

    h2(doc, "4.4  The Sidebar Navigation")
    body(doc,
         "The left sidebar lists ten navigation buttons: seven active features and "
         "three Coming Soon placeholders. The navigation logic in JavaScript works like this:")
    code_block(doc,
               "// When any active nav button is clicked:\n"
               "document.querySelectorAll('.nav-btn:not(.nav-btn-soon)').forEach(btn => {\n"
               "    btn.addEventListener('click', () => {\n"
               "        // 1. Remove 'active' class from all buttons\n"
               "        // 2. Add 'active' class to clicked button\n"
               "        // 3. Hide all content panels\n"
               "        // 4. Show the panel matching the button's data-tab attribute\n"
               "    });\n"
               "});")
    body(doc,
         "The :not(.nav-btn-soon) selector ensures that Coming Soon buttons are "
         "excluded from this click handler. Coming Soon buttons still show a panel "
         "when clicked (describing the planned feature), but they don't get styled "
         "as 'active' and don't trigger any API calls.")

    h2(doc, "4.5  The AI Chat Tab — Feature by Feature")

    h3(doc, "Quick Prompt Buttons")
    body(doc,
         "Buttons like 'Explain Phishing', 'Password Tips', 'Network Security 101' "
         "pre-fill the chat input with common questions. This lowers the barrier for "
         "users who don't know what to ask, and demonstrates the system's "
         "capabilities at a glance.")

    h3(doc, "The Typing Indicator")
    body(doc,
         "When a message is sent, the code shows animated 'typing dots' (three dots "
         "that pulse in sequence) to indicate the AI is thinking. This is important "
         "because AI responses can take 5-20 seconds to arrive — without this "
         "indicator, users might think the system has crashed. The indicator is removed "
         "when the response arrives.")

    h3(doc, "Text Formatting (formatAIText)")
    body(doc,
         "The AI model formats its responses using Markdown conventions — **text** "
         "for bold, bullet points starting with '-' or '*', double newlines for "
        "paragraphs. Since the chat displays HTML, the formatAIText() function converts "
         "these to proper HTML: **text** → <strong>text</strong>, newlines → <br> tags. "
         "This makes responses much more readable than raw text.")

    h3(doc, "NETWORK_ERROR Detection")
    body(doc,
         "When a response starts with 'NETWORK_ERROR:', the JavaScript code "
         "recognises this prefix and triggers two actions: (1) displays the error "
         "message with a special red 'Offline' formatting, and (2) automatically "
         "re-runs the connection check to update the status badge. This means the "
         "user gets immediate feedback and the status indicator updates without "
         "requiring any extra action.")

    h2(doc, "4.6  The Password Strength Analyser — How Client-Side Works")
    body(doc,
         "This is the only feature that is entirely client-side, meaning it runs "
         "100% in the browser with no communication to the server. Here is why "
         "this matters: your password is never sent anywhere. Not to the Flask "
         "server, not to Hugging Face, nowhere. It stays in your browser. This "
         "is a fundamental privacy and security principle.")

    h3(doc, "Entropy Calculation")
    body(doc,
         "Entropy is a mathematical measure of unpredictability. The more possible "
         "characters and the longer the password, the higher the entropy, the harder "
         "to crack.")
    code_block(doc,
               "function getEntropy(password) {\n"
               "    let charSet = 0;\n"
               "    if (/[a-z]/.test(password)) charSet += 26;  // lowercase letters\n"
               "    if (/[A-Z]/.test(password)) charSet += 26;  // uppercase letters\n"
               "    if (/[0-9]/.test(password)) charSet += 10;  // digits\n"
               "    if (/[^a-zA-Z0-9]/.test(password)) charSet += 32; // symbols\n"
               "    return password.length * Math.log2(charSet || 1);\n"
               "}")
    body(doc,
         "For example: 'Pa$$w0rd' uses uppercase, lowercase, digits, and symbols "
         "(charSet = 94). Length = 8. Entropy = 8 × log2(94) ≈ 52 bits. "
         "At 10 billion guesses per second (a high-end GPU), cracking a 52-bit "
         "password would take about 12 days. A 90-bit password would take "
         "39 billion years — effectively uncrackable.")

    h3(doc, "The Six Criteria Checks")
    bul(doc, [
        "At least 8 characters",
        "At least 12 characters",
        "Contains uppercase letters",
        "Contains lowercase letters",
        "Contains numbers",
        "Contains special characters (!, @, #, $ etc.)",
    ])
    body(doc,
         "Each criterion is checked with a simple regular expression test "
         "(e.g., /[A-Z]/.test(password)) and immediately updates a visual "
         "checklist with green ticks or red crosses as the user types.")

    h2(doc, "4.7  The Security Quiz — How Scoring Works")
    body(doc,
         "When a question is fetched from /api/quiz, the JSON response contains "
         "the correct answer letter (A, B, C, or D). When the user clicks an option "
         "button, JavaScript compares the clicked letter to the correct answer letter:")
    bul(doc, [
        "If correct: the clicked button turns green, score increments by 1, "
        "total questions increments by 1.",
        "If incorrect: the clicked button turns red, the correct button turns green "
        "(so user can see the right answer), score stays the same, total questions "
        "increments by 1.",
        "In both cases: the explanation text from the JSON is displayed below the "
        "options, and a 'Next Question' button appears.",
    ])
    body(doc,
         "The score persists across the session (as long as the page is open) "
         "and is displayed as 'X / Y correct' where X is score and Y is total "
         "questions attempted.")

    h2(doc, "4.8  Coming Soon Features")
    body(doc,
         "Three features in the sidebar are marked 'Coming Soon' with an orange "
         "badge: Threat Scanner, Dark Web Monitor, and File Analyzer. Each of these "
         "has a dedicated information panel that explains what the feature will do "
         "when implemented.")

    h3(doc, "Why 'Coming Soon' and not 'Not Available'?")
    body(doc,
         "These features were designed as Coming Soon (rather than simply removed) "
         "for several important reasons:")
    bul(doc, [
        "They represent genuine planned functionality that is architecturally "
        "reasonable but out of scope for the current FYP phase.",
        "Showing them demonstrates forward-thinking system design — the examiner "
        "can see that the system was designed with extensibility in mind.",
        "They are technically out of scope for a final year project because they "
        "require: live network scanning (Nmap/Shodan APIs), access to dark web "
        "APIs (which have ethical/legal considerations), and malware sandboxing "
        "infrastructure (VirusTotal/Cuckoo).",
        "The Coming Soon design pattern is standard in commercial software products "
        "to communicate roadmap features.",
    ])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5 — THE AI INTEGRATION EXPLAINED
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 5: THE AI INTEGRATION — HOW THE INTELLIGENCE WORKS")

    h2(doc, "5.1  What is Hugging Face?")
    body(doc,
         "Hugging Face is the world's largest platform for sharing, discovering, and "
         "running AI models. Think of it as GitHub, but for AI models instead of "
         "code. It hosts thousands of open-source AI models including Meta's LLaMA "
         "family. The 'Inference API' is a service Hugging Face provides that lets "
         "you run these models via HTTP requests — you send a message, they run the "
         "model, they send back the response. You do not need a powerful computer "
         "to use it; all the AI computation happens on Hugging Face's servers.")

    h2(doc, "5.2  What is Meta Llama 3.1-8B-Instruct?")
    body(doc,
         "Let's break down the name:")
    bul(doc, [
        "Meta — the company that created and released this AI model (Meta AI, "
        "the research arm of Facebook/Meta Platforms).",
        "Llama — the name of the model family (Large Language Model Meta AI).",
        "3.1 — the version number (LLaMA 3.1 was released in 2024).",
        "8B — 8 billion parameters. Parameters are the 'learned values' inside "
        "the neural network. More parameters generally means more capable, but "
        "also requires more computing power. 8B is a 'small-large' model — powerful "
        "enough for most advisory tasks but fast and affordable to run.",
        "Instruct — this variant has been fine-tuned specifically to follow "
        "instructions. The base LLaMA model predicts text; the Instruct version "
        "has been trained with human feedback to follow commands and answer questions "
        "helpfully and safely.",
    ])

    h2(doc, "5.3  The OpenAI-Compatible Format — Why It Matters")
    body(doc,
         "The Hugging Face Inference API, for chat models, uses the same message "
         "format that OpenAI's GPT API uses. This is deliberately done for "
         "compatibility. The format structures a conversation as a list of messages, "
         "each with a 'role' (who is speaking) and 'content' (what they said):")
    code_block(doc,
               'messages = [\n'
               '    {"role": "system",    "content": "You are a cybersecurity expert..."},\n'
               '    {"role": "user",      "content": "What is phishing?"},\n'
               '    {"role": "assistant", "content": "Phishing is..."},\n'
               '    {"role": "user",      "content": "How do I avoid it?"}\n'
               ']')
    body(doc,
         "Roles:")
    bul(doc, [
        "system — background instructions for the AI (our SYSTEM_PROMPT). Sent "
        "once at the start of every conversation.",
        "user — messages from the human user.",
        "assistant — the AI's previous responses (included for context in "
        "multi-turn conversations).",
    ])
    body(doc,
         "This format means that if we ever wanted to switch from LLaMA to a "
         "different model (like Mistral or GPT-4), we would only need to change "
         "the model name and API URL — the rest of the code works identically. "
         "This is the power of using an industry-standard API format.")

    h2(doc, "5.4  Why Was the 8B Model Chosen (Not 70B or 405B)?")
    bul(doc, [
        "Free tier access: Hugging Face's free Inference API supports the 8B model. "
        "Larger models require paid plans.",
        "Speed: The 8B model responds faster, which is important for an interactive "
        "advisory tool where users expect quick responses.",
        "Sufficient capability: For the cybersecurity advisory tasks in this project "
        "(explanation, analysis, quiz generation, tips), the 8B model produces "
        "high-quality, accurate responses that meet all functional requirements.",
        "Academic scope: Using the 8B model is entirely appropriate for a final year "
        "project. Larger models provide marginal improvements for these tasks but "
        "would require commercial API access.",
    ])

    h2(doc, "5.5  Context Windows and History Limits")
    body(doc,
         "LLMs have a 'context window' — the maximum amount of text they can process "
         "in a single API call. For Llama 3.1-8B, this is 128,000 tokens (roughly "
         "96,000 words). Our conversation history is limited to 12 messages not "
         "because the model can't handle more, but as a practical constraint: "
         "longer prompts take longer to process, cost more tokens, and in a real "
         "deployment would increase API costs. 12 messages is sufficient to maintain "
         "meaningful conversational context for typical advisory sessions.")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6 — HOW TO EXPLAIN YOUR PROJECT
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 6: HOW TO EXPLAIN YOUR PROJECT — EXAMINER Q&A PREPARATION")

    h2(doc, "6.1  The 60-Second Elevator Pitch")
    body(doc,
         "If your examiner asks 'what did you build?' — this is the answer:")
    body(doc,
         "\"I built CyberGuard AI — a web-based cybersecurity advisory system powered "
         "by Meta's Llama 3.1 AI model. A user opens the website in their browser "
         "and can have a conversation with an AI cybersecurity expert, explore a "
         "library of twelve cyber threats, take a security awareness quiz, check "
         "their password strength, analyse suspicious URLs for phishing, look up "
         "software vulnerabilities in plain English, and get daily security tips. "
         "The backend is Python Flask, the frontend is HTML, CSS, and JavaScript, "
         "and the AI is accessed through the Hugging Face Inference API.\"")

    h2(doc, "6.2  Common Examiner Questions and Model Answers")

    qa(doc,
       "Why did you choose Flask over Django?",
       "Flask is a micro-framework — it gives you only what you need with minimal "
       "configuration. Django includes an ORM, admin panel, and many features I "
       "didn't need for this project. Flask was the right tool because my project "
       "is API-centric with a static frontend, not a database-driven web app. "
       "Flask also works perfectly with Python 3.6 which was the version I had available.")

    qa(doc,
       "Why did you use Hugging Face instead of OpenAI's API?",
       "Hugging Face provides free access to open-source models including Meta's LLaMA, "
       "which made it ideal for an academic project with no budget for commercial API "
       "access. LLaMA 3.1-8B-Instruct is powerful enough for all the advisory tasks "
       "in this project. Additionally, using open-source models aligns with academic "
       "research values — the model's weights and training methodology are published "
       "and auditable, unlike proprietary models.")

    qa(doc,
       "How does the AI know to be a cybersecurity expert and not just a general chatbot?",
       "This is achieved through the system prompt — a set of instructions included "
       "at the beginning of every API call. The system prompt tells the AI: its name "
       "is CyberGuard AI, it is a cybersecurity expert, it should focus on defensive "
       "measures, and it must never help with actual attacks. This is called prompt "
       "engineering. The AI model itself is general-purpose, but the system prompt "
       "constrains and directs its behaviour for our specific use case.")

    qa(doc,
       "What happens if the internet is not available?",
       "The system handles this gracefully. The call_llama() function catches "
       "ConnectionError exceptions and returns a message prefixed with 'NETWORK_ERROR:'. "
       "The frontend JavaScript detects this prefix and shows the user a clear error "
       "message with four troubleshooting steps. Additionally, a connection status "
       "endpoint (/api/status) is checked on page load, and a red banner is displayed "
       "at the top of the page when the API is unreachable. The password analyser "
       "continues to work even offline because it is entirely client-side.")

    qa(doc,
       "How does the password analyser work without sending the password to the server?",
       "The password is never sent to the server — all calculations happen in the "
       "user's browser using JavaScript. The script calculates Shannon entropy using "
       "the formula: entropy = length × log₂(character_set_size). It also runs six "
       "regular expression checks for character type criteria. Everything happens "
       "instantly on the client side as the user types. This is a privacy-preserving "
       "design — a password should never be transmitted over a network just to be "
       "evaluated.")

    qa(doc,
       "Why did you mark some features 'Coming Soon' instead of implementing them?",
       "The three Coming Soon features — live network threat scanner, dark web monitor, "
       "and file malware analyser — require infrastructure and capabilities beyond "
       "the scope of a final year project. Network scanning requires Nmap integration "
       "and appropriate network permissions. Dark web monitoring requires Tor API "
       "access and raises ethical and legal considerations. File malware analysis "
       "requires a sandboxing environment like Cuckoo Sandbox or VirusTotal API. "
       "I scoped them as Coming Soon to demonstrate architectural thinking and "
       "system design awareness while keeping the project deliverable within the "
       "available time and resources.")

    qa(doc,
       "What is the purpose of the .env file?",
       "The .env file stores sensitive configuration values — specifically the Hugging "
       "Face API key and Flask secret key — separately from the source code. If the "
       "API key were hardcoded in app.py and the code was shared publicly on GitHub, "
       "anyone who found it could use my API quota. The .env file is excluded from "
       "version control (added to .gitignore), so the API key stays private. The "
       "python-dotenv library reads this file and makes the values available as "
       "environment variables that app.py accesses via os.getenv().")

    qa(doc,
       "What is CORS and why do you need it?",
       "CORS stands for Cross-Origin Resource Sharing. Browsers have a security policy "
       "that prevents a web page from making API requests to a different 'origin' "
       "(domain, port, or protocol) than the one it was loaded from. Flask runs on "
       "port 5000; without CORS configuration, the browser would block the frontend "
       "JavaScript from calling /api/chat. Flask-CORS adds the necessary HTTP headers "
       "that tell the browser it is safe to make these cross-origin requests.")

    qa(doc,
       "How does the quiz ensure the AI returns valid JSON?",
       "Two mechanisms work together. First, the prompt explicitly instructs the AI "
       "to return ONLY a JSON object in a specific schema — no extra text. Second, "
       "even with this instruction, AI models sometimes add explanatory text around "
       "the JSON. The code uses a regular expression (re.search(r'\\{[\\s\\S]*\\}', raw)) "
       "to find and extract just the JSON object from whatever the AI returns. If "
       "parsing still fails, the endpoint returns {\"raw\": raw, \"parse_error\": true} "
       "so the frontend can handle the edge case gracefully.")

    qa(doc,
       "What is temperature in the context of the AI model?",
       "Temperature is a parameter that controls how random or 'creative' the AI's "
       "responses are. At temperature 0, the model always picks the most statistically "
       "probable next token — responses are deterministic and repetitive. At temperature "
       "1.0, it picks more randomly — creative but potentially inaccurate. We use 0.7 "
       "as a balance: responses are varied enough to feel natural, but grounded "
       "enough to be accurate for technical cybersecurity topics.")

    qa(doc,
       "How does multi-turn conversation work?",
       "Each API call to /api/chat includes the full conversation history as an array "
       "of messages. The client's JavaScript maintains this history array locally and "
       "sends the last 12 messages with every new request. The backend prepends the "
       "system prompt and appends the new user message, then sends the complete array "
       "to the Hugging Face API. The AI model processes all of this context together "
       "and generates a response that is aware of the entire conversation history, "
       "creating the appearance of memory.")

    qa(doc,
       "What is the difference between your URL analyser and VirusTotal?",
       "VirusTotal sends the URL to 70+ antivirus engines and reputation services "
       "that actually visit and analyse the URL. My URL analyser uses AI to perform "
       "structural analysis only — it analyses the text of the URL itself for "
       "suspicious patterns without fetching it. VirusTotal is more comprehensive "
       "for definitive threat detection; my analyser is educational and explains "
       "WHY a URL looks suspicious, which is more valuable for teaching users to "
       "recognise phishing on their own. My system also does not create a record "
       "of submitted URLs on a third-party platform.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7 — THE TECHNOLOGY STACK FULLY EXPLAINED
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 7: THE COMPLETE TECHNOLOGY STACK — WHAT EACH TOOL IS AND WHY")

    h2(doc, "7.1  Python 3.6")
    body(doc,
         "Python is the programming language used for the backend. Python was chosen "
         "because it is the dominant language for AI and data science work, has "
         "excellent libraries for HTTP requests (requests), web servers (Flask), and "
         "environment management (python-dotenv). Python 3.6 was used because it was "
         "the version available on the development machine. This required downgrading "
         "some libraries (Flask to version 2.0.3 instead of the latest 3.x, which "
         "requires Python 3.8+) to ensure compatibility.")

    h2(doc, "7.2  Flask 2.0.3")
    body(doc,
         "Flask is the web framework. Version 2.0.3 was specifically chosen (not the "
         "latest 3.x) for Python 3.6 compatibility. Flask handles: routing (matching "
         "URLs to Python functions), serving static files (CSS, JS), rendering HTML "
         "templates, and returning JSON responses. It runs a development web server "
         "on port 5000 for local testing.")

    h2(doc, "7.3  requests 2.27.1")
    body(doc,
         "The requests library is Python's standard HTTP client. It is used to make "
         "HTTP POST requests from the Flask server to the Hugging Face API. The "
         "requests.Session class enables connection pooling and persistent settings "
         "(like proxy configuration) across multiple API calls.")

    h2(doc, "7.4  python-dotenv 0.19.2")
    body(doc,
         "Reads key=value pairs from the .env file and makes them available as "
         "environment variables. This is the industry-standard approach to managing "
         "application secrets (API keys, database passwords, etc.) — keeping them "
         "out of source code and version control.")

    h2(doc, "7.5  flask-cors 3.0.10")
    body(doc,
         "Adds CORS headers to Flask responses so the browser allows the frontend "
         "JavaScript to make API calls to the Flask server. One line of code "
         "(CORS(app)) enables this globally for all endpoints.")

    h2(doc, "7.6  HTML5")
    body(doc,
         "The markup language for the web page structure. Key HTML5 features used: "
         "<canvas> for the matrix rain animation, semantic elements (<main>, <aside>, "
         "<section>), the Fetch API (technically JavaScript, but relies on HTML5 "
         "browser standards), and the <audio> tag support (not used, but available).")

    h2(doc, "7.7  CSS3 with Custom Properties")
    body(doc,
         "Stylesheet language for visual presentation. Key CSS3 features used: "
         "custom properties (variables) for the colour scheme, CSS transitions for "
         "smooth animations, flexbox layout for the sidebar + main content structure, "
         "media queries (responsive design), and keyframe animations for the typing "
         "indicator dots.")

    h2(doc, "7.8  Vanilla JavaScript (ES6+)")
    body(doc,
         "'Vanilla' means plain JavaScript with no framework (no React, Vue, Angular). "
         "This was deliberately chosen to keep the project simple, reduce dependencies, "
         "and demonstrate fundamental JavaScript skills. ES6+ features used: async/await "
         "for API calls, arrow functions, template literals, destructuring, the Fetch "
         "API, and the Canvas 2D API for the animation.")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 8 — FILE-BY-FILE QUICK REFERENCE
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 8: FILE-BY-FILE QUICK REFERENCE")

    h2(doc, "Project Folder Structure")
    code_block(doc,
               "cybersec-advisor/\n"
               "├── app.py                   ← Python Flask backend (the brain)\n"
               "├── requirements.txt          ← Python package dependencies\n"
               "├── .env                      ← Secret API key (never share/commit)\n"
               "├── templates/\n"
               "│   └── index.html            ← The web page HTML structure\n"
               "└── static/\n"
               "    ├── css/\n"
               "    │   └── style.css         ← All visual styling\n"
               "    └── js/\n"
               "        └── app.js            ← All browser-side JavaScript logic")

    h2(doc, "How to Start the System")
    code_block(doc,
               "# 1. Navigate to project folder\n"
               "cd cybersec-advisor\n\n"
               "# 2. Install dependencies (first time only)\n"
               "pip install -r requirements.txt\n\n"
               "# 3. Start the Flask server\n"
               "python app.py\n\n"
               "# 4. Open browser at:\n"
               "http://127.0.0.1:5000")

    h2(doc, "What Each Function in app.js Does")
    funcs = [
        ("initMatrixRain()", "Starts the matrix animation canvas"),
        ("checkConnection()", "Calls /api/status, updates the Online/Offline badge and red banner"),
        ("sendChat()", "Sends the user's message to /api/chat, displays typing dots, shows response"),
        ("appendMessage(role, text)", "Adds a message bubble to the chat display"),
        ("formatAIText(text)", "Converts **bold** and newlines to HTML formatting"),
        ("loadThreatAnalysis(threat)", "Calls /api/threat-analysis for a clicked threat card"),
        ("loadQuiz()", "Calls /api/quiz, renders the MCQ question"),
        ("checkAnswer(selected)", "Evaluates quiz answer, updates score, shows explanation"),
        ("analyzePassword()", "Calculates entropy, updates strength bar and criteria checks (client-side)"),
        ("analyzeURL()", "Sends URL to /api/analyze-url, displays risk assessment"),
        ("lookupCVE()", "Sends CVE to /api/explain-cve, displays explanation"),
        ("loadTip()", "Calls /api/security-tip, displays the daily tip"),
        ("getEntropy(password)", "Calculates Shannon entropy for password strength"),
        ("getCrackTime(entropy)", "Converts entropy bits to human-readable crack time estimate"),
    ]
    for name, desc in funcs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _r(p, f"{name}  ", bold=True, mono=True, size=10)
        _r(p, f"— {desc}")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 9 — GLOSSARY
    # ═══════════════════════════════════════════════════════════════════════
    h1(doc, "SECTION 9: GLOSSARY — KEY TERMS EXPLAINED SIMPLY")

    terms = [
        ("API (Application Programming Interface)",
         "A way for two software programs to talk to each other. When our Flask "
         "server talks to Hugging Face, it uses their API — a defined set of "
         "rules for sending requests and receiving responses."),
        ("Bearer Token",
         "A type of API authentication. 'Bearer [key]' in the HTTP header means "
         "'the person sending this request holds this key and is authorised.'"),
        ("Context Window",
         "The maximum amount of text an AI model can process in one call. Like "
         "short-term memory — the AI can only 'see' text within this window."),
        ("CVE (Common Vulnerabilities and Exposures)",
         "A standardised naming system for software security flaws. Each CVE has "
         "a unique ID (e.g., CVE-2021-44228) that security researchers, tools, "
         "and news articles use to refer to the same vulnerability."),
        ("CVSS (Common Vulnerability Scoring System)",
         "A 0-10 score rating the severity of a software vulnerability. 9.0-10 is "
         "Critical, 7.0-8.9 is High, 4.0-6.9 is Medium, 0.1-3.9 is Low."),
        ("DOM (Document Object Model)",
         "The browser's representation of the web page as a tree of elements that "
         "JavaScript can read and modify dynamically."),
        ("Endpoint",
         "A specific URL in an API that performs a specific function. "
         "/api/chat is an endpoint that handles AI conversation requests."),
        ("Entropy (Password)",
         "A measure of how unpredictable a password is, in bits. Higher entropy "
         "means harder to crack by brute force."),
        ("Flask",
         "A Python web framework — a toolkit that makes it easy to create web "
         "servers and APIs in Python with minimal code."),
        ("Hugging Face",
         "A company and platform that hosts thousands of open-source AI models "
         "and provides APIs to run them without needing your own GPU."),
        ("JSON (JavaScript Object Notation)",
         "A text format for structured data, used for API communication. "
         "Example: {\"name\": \"Alice\", \"age\": 30}."),
        ("LLM (Large Language Model)",
         "An AI model trained on massive amounts of text that can understand and "
         "generate human language. Examples: GPT-4, LLaMA, Gemini."),
        ("Phishing",
         "A cyber attack where criminals create fake emails or websites that look "
         "legitimate to trick users into giving up passwords or personal information."),
        ("Prompt Engineering",
         "The practice of carefully designing the instructions given to an AI model "
         "to get the desired type, format, and quality of output."),
        ("Proxy Server",
         "An intermediary computer through which network traffic is routed. Common "
         "in university and corporate networks. CyberGuard AI supports proxy "
         "configuration via environment variables."),
        ("REST API",
         "A standard style of API design that uses HTTP methods (GET, POST) and "
         "URLs to provide access to resources. All CyberGuard AI's endpoints "
         "follow REST conventions."),
        ("Single-Page Application (SPA)",
         "A web app that loads one HTML page and dynamically updates content "
         "without full page reloads. CyberGuard AI is an SPA."),
        ("Temperature (AI)",
         "A parameter controlling AI response randomness. 0 = deterministic "
         "and repetitive, 1 = very random and creative. We use 0.7."),
        ("Token (AI)",
         "The basic unit of text an AI model processes — roughly 3/4 of a word. "
         "max_tokens limits the length of the AI's response."),
        ("Typosquatting",
         "A phishing technique where attackers register domain names that look "
         "like legitimate sites but with spelling errors (e.g., paypa1.com "
         "instead of paypal.com)."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(term + ":  ")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r2 = p.add_run(definition)
        r2.font.name = "Calibri"
        r2.font.size = Pt(12)

    doc.save(OUT_PATH)
    print(f"[OK] Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
